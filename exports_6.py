"""
[exports_6-001] 导出服务模块
封装报告导出功能，支持 Word、Excel 等格式

核心职责:
- ExportService: 提供 Word、Excel 导出服务
- 导出 Word 报告：包含统计表格和 Markdown 正文
- 导出 Excel 统计表和数据

为什么这样设计:
将导出逻辑封装为服务类，便于扩展新的导出格式，
同时提供统一的 API 接口。
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from typing import Dict, Any, List, Optional
import os
from datetime import datetime
import re


# ==================== 辅助函数 ====================

def set_cell_shading(cell, color: str) -> None:
    """
    [exports_6-001-01] 设置单元格背景色
    Word 表格单元格背景色设置

    Args:
        cell: Word 表格单元格
        color: 十六进制颜色值（不含#）
    """
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_markdown_to_doc(doc: Document, markdown_text: str) -> None:
    """
    [exports_6-001-02] 添加 Markdown 到 Word
    将 Markdown 文本解析并添加到 Word 文档

    核心职责:
    - 解析标题（###, ##, #）
    - 解析列表（- 或 1.）
    - 解析普通段落

    为什么这样设计:
    AI 生成的报告是 Markdown 格式，需要转换为 Word 格式
    """
    lines = markdown_text.split('\n')

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # 标题
        if stripped.startswith('### '):
            heading = doc.add_heading(stripped[4:], 3)
            for run in heading.runs:
                run.font.size = Pt(14)
        elif stripped.startswith('## '):
            heading = doc.add_heading(stripped[3:], 2)
            for run in heading.runs:
                run.font.size = Pt(16)
        elif stripped.startswith('# '):
            heading = doc.add_heading(stripped[2:], 1)
            for run in heading.runs:
                run.font.size = Pt(18)

        # 列表项
        elif stripped.startswith('- ') or stripped.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, stripped[2:])

        # 编号列表
        elif re.match(r'^\d+\.\s', stripped):
            para = doc.add_paragraph(style='List Number')
            _add_formatted_text(para, re.sub(r'^\d+\.\s', '', stripped))

        # 普通段落
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)


def _add_formatted_text(para, text: str) -> None:
    """
    [exports_6-001-03] 添加带格式的文本
    处理 Markdown 内联格式（加粗、代码）

    Args:
        para: Word 段落对象
        text: 文本内容
    """
    # 处理加粗 **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            # 处理行内代码 `code`
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = para.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    para.add_run(cp)


# ==================== 导出服务类 ====================

class ExportService:
    """[exports_6-001-10] 导出服务
    提供报告导出功能

    核心职责:
    - export_to_word: 导出 Word 报告
    - export_to_excel: 导出 Excel 表格
    - export_full_report: 导出完整报告包

    为什么这样设计:
    封装导出逻辑，提供统一的 API，便于上层调用
    """

    def __init__(self, output_dir: str = "temp/exported"):
        """
        初始化导出服务

        Args:
            output_dir: 默认输出目录
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def export_to_word(
        self,
        report_markdown: str,
        stats_result: Dict,
        is_comparison: bool = False,
        project_name: str = "标注评测报告",
        output_path: Optional[str] = None
    ) -> str:
        """
        [exports_6-001-11] 导出 Word 报告
        将评测报告导出为 Word 文档

        核心职责:
        - 创建文档结构（标题、统计概览、表格、正文）
        - 处理对比模式和普通模式
        - 解析 Markdown 正文

        Args:
            report_markdown: Markdown 格式的报告正文
            stats_result: 统计结果字典
            is_comparison: 是否为对比模式
            project_name: 项目名称
            output_path: 输出路径（可选）

        Returns:
            str: 导出文件路径
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(11)

        # 标题
        title = doc.add_heading(project_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1d, 0x6f, 0x64)

        # 生成时间
        time_para = doc.add_paragraph()
        time_run = time_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 统计概览
        doc.add_heading("统计概览", 1)

        if is_comparison:
            # 对比模式
            para = doc.add_paragraph()
            para.add_run("对比字段：").bold = True
            para.add_run(f"{stats_result.get('compare_field', '-')}")

            para = doc.add_paragraph()
            para.add_run(f"版本 A ({stats_result.get('version_a', '-')}): ").bold = True
            para.add_run(f"{stats_result.get('denominator_count_a', 0)} 条")

            para = doc.add_paragraph()
            para.add_run(f"版本 B ({stats_result.get('version_b', '-')}): ").bold = True
            para.add_run(f"{stats_result.get('denominator_count_b', 0)} 条")
        else:
            para = doc.add_paragraph()
            para.add_run("公共分母：").bold = True
            para.add_run(f"{stats_result.get('denominator_count', 0)} 条")

            if stats_result.get('denominator_description'):
                para = doc.add_paragraph()
                para.add_run("分母说明：").bold = True
                para.add_run(stats_result['denominator_description'])

        doc.add_paragraph()

        # 统计表格
        doc.add_heading("指标统计表", 1)

        results = stats_result.get("results", [])

        if is_comparison:
            # 对比模式表格
            table = doc.add_table(rows=len(results) + 1, cols=6)
            table.style = 'Table Grid'

            # 表头
            headers = ["指标名称", f"版本 A ({stats_result.get('version_a', '-')})", "",
                       f"版本 B ({stats_result.get('version_b', '-')})", "", "Gap"]
            sub_headers = ["", "分子/分母", "百分比", "分子/分母", "百分比", ""]

            header_row = table.rows[0]
            for i, h in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = h
                set_cell_shading(cell, "1d6f64")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)

            # 数据行
            for idx, r in enumerate(results):
                row = table.rows[idx + 1]

                row.cells[0].text = r.get("name", "-")

                row.cells[1].text = f"{r.get('numerator_a', 0)}/{r.get('denominator_a', 0)}"
                pct_a = r.get('percentage_a')
                row.cells[2].text = f"{pct_a}%" if pct_a is not None else "N/A"

                row.cells[3].text = f"{r.get('numerator_b', 0)}/{r.get('denominator_b', 0)}"
                pct_b = r.get('percentage_b')
                row.cells[4].text = f"{pct_b}%" if pct_b is not None else "N/A"

                gap = r.get("gap")
                trend = r.get("trend", "")
                if gap is not None:
                    trend_symbol = "↑" if trend == "improved" else ("↓" if trend == "declined" else "→")
                    row.cells[5].text = f"{gap}% {trend_symbol}"
                else:
                    row.cells[5].text = "N/A"

        else:
            # 普通模式表格
            table = doc.add_table(rows=len(results) + 1, cols=4)
            table.style = 'Table Grid'

            # 表头
            headers = ["指标名称", "分子", "分母", "百分比"]
            header_row = table.rows[0]
            for i, h in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = h
                set_cell_shading(cell, "1d6f64")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)

            # 数据行
            for idx, r in enumerate(results):
                row = table.rows[idx + 1]

                row.cells[0].text = r.get("name", "-")
                row.cells[1].text = str(r.get("numerator", 0))
                row.cells[2].text = str(r.get("denominator", 0))

                pct = r.get("percentage")
                row.cells[3].text = f"{pct}%" if pct is not None else "N/A"

        doc.add_paragraph()

        # 报告正文
        doc.add_heading("评测报告", 1)

        # 解析 Markdown 并添加到文档
        _add_markdown_to_doc(doc, report_markdown)

        # 保存文档
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"评测报告_{timestamp}.docx"
            output_path = os.path.join(self.output_dir, filename)

        doc.save(output_path)
        return output_path

    def export_indicators_to_excel(
        self,
        indicators_data: pd.DataFrame,
        project_name: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        [exports_6-001-12] 导出指标 Excel
        将指标数据导出为 Excel 文件

        Args:
            indicators_data: 指标数据 DataFrame
            project_name: 项目名称
            output_path: 输出路径（可选）

        Returns:
            str: 导出文件路径
        """
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{project_name}_{timestamp}.xlsx"
            output_path = os.path.join(self.output_dir, filename)

        indicators_data.to_excel(output_path, index=False, engine='openpyxl')
        return output_path

    def export_filtered_cases_to_excel(
        self,
        df: pd.DataFrame,
        project_name: str,
        filter_name: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        [exports_6-001-13] 导出筛选 Case
        将筛选后的案例导出为 Excel

        Args:
            df: 筛选后的 DataFrame
            project_name: 项目名称
            filter_name: 筛选条件名称
            output_path: 输出路径（可选）

        Returns:
            str: 导出文件路径
        """
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{project_name}_{filter_name}_case_{timestamp}.xlsx"
            output_path = os.path.join(self.output_dir, filename)

        df.to_excel(output_path, index=False, engine='openpyxl')
        return output_path

    def export_full_report(
        self,
        project_name: str,
        stats_result: Dict,
        metrics_config: Dict,
        df: pd.DataFrame,
        report_markdown: str = ""
    ) -> Dict[str, str]:
        """
        [exports_6-001-14] 导出完整报告包
        一次性导出所有相关文件

        核心职责:
        - Word 报告
        - 统计 Excel 表
        - 口径配置 JSON
        - 原始数据 Excel

        Args:
            project_name: 项目名称
            stats_result: 统计结果
            metrics_config: 口径配置
            df: 原始数据
            report_markdown: 报告正文（可选）

        Returns:
            Dict[str, str]: 各导出文件的路径
        """
        os.makedirs(self.output_dir, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        result = {}

        # 1. 导出 Word 报告
        if report_markdown:
            is_comparison = "version_a" in stats_result
            word_path = os.path.join(self.output_dir, f"{project_name}_评测报告_{timestamp}.docx")
            self.export_to_word(report_markdown, stats_result, is_comparison, project_name, word_path)
            result["report"] = word_path

        # 2. 导出统计表
        stats_df = self._format_stats_for_excel(stats_result)
        stats_path = os.path.join(self.output_dir, f"{project_name}_统计表_{timestamp}.xlsx")
        stats_df.to_excel(stats_path, index=False, engine='openpyxl')
        result["stats"] = stats_path

        # 3. 导出口径配置
        import json
        config_path = os.path.join(self.output_dir, f"{project_name}_口径配置_{timestamp}.json")
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(metrics_config, f, ensure_ascii=False, indent=2)
        result["config"] = config_path

        # 4. 导出原始数据
        data_path = os.path.join(self.output_dir, f"{project_name}_原始数据_{timestamp}.xlsx")
        df.to_excel(data_path, index=False, engine='openpyxl')
        result["data"] = data_path

        return result

    def _format_stats_for_excel(self, stats_result: Dict) -> pd.DataFrame:
        """
        [exports_6-001-15] 格式化统计结果为 DataFrame
        将统计结果转换为适合 Excel 导出的格式

        Args:
            stats_result: 统计结果字典

        Returns:
            pd.DataFrame: 格式化后的表格
        """
        is_comparison = "version_a" in stats_result
        is_grouping = "group_results" in stats_result or "groups" in stats_result

        # 分组模式
        if is_grouping:
            rows = []
            group_results = stats_result.get("group_results", [])
            for gr in group_results:
                group_name = gr.get("group_name", gr.get("group_value", ""))
                for r in gr.get("results", []):
                    pct = r.get("percentage")
                    rows.append({
                        "分组": group_name,
                        "指标名称": r.get("name", ""),
                        "分子": r.get("numerator", 0),
                        "分母": r.get("denominator", 0),
                        "百分比": f"{pct}%" if pct is not None else "N/A"
                    })
            return pd.DataFrame(rows)

        if is_comparison:
            rows = []
            for r in stats_result.get("results", []):
                gap = r.get("gap")
                trend = r.get("trend", "")

                trend_symbol = ""
                winner = ""
                if trend == "improved":
                    trend_symbol = "↑"
                    winner = "右边好"
                elif trend == "declined":
                    trend_symbol = "↓"
                    winner = "左边好"
                elif trend == "unchanged":
                    trend_symbol = "→"
                    winner = "平局"
                else:
                    winner = "-"

                rows.append({
                    "指标名称": r.get("name", ""),
                    "版本 A 分子": r.get("numerator_a", 0),
                    "版本 A 分母": r.get("denominator_a", 0),
                    "版本 A(%)": r.get("percentage_a"),
                    "版本 B 分子": r.get("numerator_b", 0),
                    "版本 B 分母": r.get("denominator_b", 0),
                    "版本 B(%)": r.get("percentage_b"),
                    "Gap(%)": gap,
                    "趋势": trend_symbol,
                    "胜者": winner
                })

            return pd.DataFrame(rows)

        # 普通模式
        rows = []
        for r in stats_result.get("results", []):
            pct = r.get("percentage")
            rows.append({
                "指标名称": r.get("name", ""),
                "分子": r.get("numerator", 0),
                "分母": r.get("denominator", 0),
                "百分比": f"{pct}%" if pct is not None else "N/A",
                "原始百分比": pct
            })

        return pd.DataFrame(rows)

    def export_all_to_zip(
        self,
        stats_result: Dict,
        report_markdown: str,
        badcase_df: Optional[pd.DataFrame],
        goodcase_df: Optional[pd.DataFrame],
        is_comparison: bool = False,
        project_name: str = "标注评测报告"
    ) -> bytes:
        """
        [exports_6-001-16] 一键导出完整报告包
        将统计结果、评测报告、Case 筛选结果打包为 ZIP

        Args:
            stats_result: 统计结果字典
            report_markdown: Markdown 格式的评测报告
            badcase_df: Bad Case 筛选结果（可选）
            goodcase_df: Good Case 筛选结果（可选）
            is_comparison: 是否对比模式
            project_name: 项目名称

        Returns:
            bytes: ZIP 文件的二进制数据
        """
        import zipfile
        import io

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # 1. 评测报告 Word
            if report_markdown:
                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'Microsoft YaHei'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                style.font.size = Pt(11)

                # 标题
                title = doc.add_heading(project_name, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in title.runs:
                    run.font.size = Pt(22)
                    run.font.color.rgb = RGBColor(0x1d, 0x6f, 0x64)

                # 生成时间
                time_para = doc.add_paragraph()
                time_run = time_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                time_run.font.size = Pt(10)
                time_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
                time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

                # 统计概览
                doc.add_heading("统计概览", 1)
                if is_comparison:
                    para = doc.add_paragraph()
                    para.add_run("对比字段：").bold = True
                    para.add_run(f"{stats_result.get('compare_field', '-')}")
                    para = doc.add_paragraph()
                    para.add_run(f"版本 A ({stats_result.get('version_a', '-')}): ").bold = True
                    para.add_run(f"{stats_result.get('denominator_count_a', 0)} 条")
                    para = doc.add_paragraph()
                    para.add_run(f"版本 B ({stats_result.get('version_b', '-')}): ").bold = True
                    para.add_run(f"{stats_result.get('denominator_count_b', 0)} 条")
                else:
                    para = doc.add_paragraph()
                    para.add_run("公共分母：").bold = True
                    para.add_run(f"{stats_result.get('denominator_count', 0)} 条")

                doc.add_paragraph()
                doc.add_heading("指标统计表", 1)

                # 统计表格
                results = stats_result.get("results", [])
                if is_comparison:
                    table = doc.add_table(rows=len(results) + 1, cols=7)
                    table.style = 'Table Grid'
                    headers = ["指标名称", "版本A分子", "版本A分母", "版本A%", "版本B分子", "版本B分母", "版本B%"]
                    for i, h in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = h
                        set_cell_shading(cell, "1d6f64")
                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.bold = True
                    for row_idx, r in enumerate(results):
                        row = table.rows[row_idx + 1]
                        row.cells[0].text = str(r.get("name", ""))
                        row.cells[1].text = str(r.get("numerator_a", 0))
                        row.cells[2].text = str(r.get("denominator_a", 0))
                        row.cells[3].text = f"{r.get('percentage_a')}%" if r.get('percentage_a') is not None else "N/A"
                        row.cells[4].text = str(r.get("numerator_b", 0))
                        row.cells[5].text = str(r.get("denominator_b", 0))
                        row.cells[6].text = f"{r.get('percentage_b')}%" if r.get('percentage_b') is not None else "N/A"
                else:
                    table = doc.add_table(rows=len(results) + 1, cols=4)
                    table.style = 'Table Grid'
                    headers = ["指标名称", "分子", "分母", "百分比"]
                    for i, h in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = h
                        set_cell_shading(cell, "1d6f64")
                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.bold = True
                    for row_idx, r in enumerate(results):
                        row = table.rows[row_idx + 1]
                        pct = r.get("percentage")
                        row.cells[0].text = str(r.get("name", ""))
                        row.cells[1].text = str(r.get("numerator", 0))
                        row.cells[2].text = str(r.get("denominator", 0))
                        row.cells[3].text = f"{pct}%" if pct is not None else "N/A"

                doc.add_paragraph()
                doc.add_heading("评测报告", 1)
                _add_markdown_to_doc(doc, report_markdown)

                # 保存到 ZIP
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                zf.writestr(f"评测报告_{timestamp}.docx", doc_buffer.read())

            # 2. 统计结果 Excel
            stats_df = self._format_stats_for_excel(stats_result)
            excel_buffer = io.BytesIO()
            stats_df.to_excel(excel_buffer, index=False, engine='openpyxl')
            excel_buffer.seek(0)
            zf.writestr(f"统计结果_{timestamp}.xlsx", excel_buffer.read())

            # 3. Bad Case Excel
            if badcase_df is not None and not badcase_df.empty:
                badcase_buffer = io.BytesIO()
                badcase_df.to_excel(badcase_buffer, index=False, engine='openpyxl')
                badcase_buffer.seek(0)
                zf.writestr(f"BadCase_{timestamp}.xlsx", badcase_buffer.read())

            # 4. Good Case Excel
            if goodcase_df is not None and not goodcase_df.empty:
                goodcase_buffer = io.BytesIO()
                goodcase_df.to_excel(goodcase_buffer, index=False, engine='openpyxl')
                goodcase_buffer.seek(0)
                zf.writestr(f"GoodCase_{timestamp}.xlsx", goodcase_buffer.read())

        zip_buffer.seek(0)
        return zip_buffer.read()

    def generate_full_markdown(
        self,
        stats_result: Dict,
        report_markdown: str,
        badcase_df: Optional[pd.DataFrame],
        goodcase_df: Optional[pd.DataFrame],
        is_comparison: bool = False
    ) -> str:
        """
        [exports_6-001-17] 生成完整 Markdown 文档
        包含统计表格、评测报告、Case 筛选结果
        可直接复制到钉钉文档

        Args:
            stats_result: 统计结果字典
            report_markdown: 评测报告正文
            badcase_df: Bad Case 数据（可选）
            goodcase_df: Good Case 数据（可选）
            is_comparison: 是否对比模式

        Returns:
            str: 完整的 Markdown 文本
        """
        lines = []
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        # 标题
        lines.append("# 标注评测报告")
        lines.append(f"\n> 生成时间：{timestamp}\n")

        # 统计概览
        lines.append("## 一、统计概览\n")
        if is_comparison:
            lines.append(f"- **对比字段**：{stats_result.get('compare_field', '-')}")
            lines.append(f"- **版本 A**（{stats_result.get('version_a', '-')}）：{stats_result.get('denominator_count_a', 0)} 条")
            lines.append(f"- **版本 B**（{stats_result.get('version_b', '-')}）：{stats_result.get('denominator_count_b', 0)} 条")
        else:
            lines.append(f"- **公共分母**：{stats_result.get('denominator_count', 0)} 条")
        lines.append("")

        # 统计表格
        lines.append("## 二、指标统计表\n")
        results = stats_result.get("results", [])

        if is_comparison:
            lines.append("| 指标名称 | 版本A分子 | 版本A分母 | 版本A% | 版本B分子 | 版本B分母 | 版本B% | Gap |")
            lines.append("|----------|-----------|-----------|--------|-----------|-----------|--------|------|")
            for r in results:
                gap = r.get("gap")
                gap_str = f"{gap}%" if gap is not None else "-"
                lines.append(f"| {r.get('name', '')} | {r.get('numerator_a', 0)} | {r.get('denominator_a', 0)} | {r.get('percentage_a', 'N/A')}% | {r.get('numerator_b', 0)} | {r.get('denominator_b', 0)} | {r.get('percentage_b', 'N/A')}% | {gap_str} |")
        else:
            lines.append("| 指标名称 | 分子 | 分母 | 百分比 |")
            lines.append("|----------|------|------|--------|")
            for r in results:
                pct = r.get("percentage")
                if pct is not None:
                    lines.append(f"| {r.get('name', '')} | {r.get('numerator', 0)} | {r.get('denominator', 0)} | {pct}% |")
                else:
                    lines.append(f"| {r.get('name', '')} | {r.get('numerator', 0)} | {r.get('denominator', 0)} | N/A |")
        lines.append("")

        # 评测报告正文
        if report_markdown:
            lines.append("## 三、评测报告分析\n")
            lines.append(report_markdown)
            lines.append("")

        # Case 筛选结果
        has_case = (badcase_df is not None and not badcase_df.empty) or (goodcase_df is not None and not goodcase_df.empty)
        if has_case:
            lines.append("## 四、Case 筛选结果\n")

            if badcase_df is not None and not badcase_df.empty:
                lines.append(f"### 问题案例（Bad Case）\n")
                lines.append(f"> 共 {len(badcase_df)} 条\n")
                preview = badcase_df.head(10)
                cols = list(preview.columns)[:6]
                lines.append("| " + " | ".join(cols) + " |")
                lines.append("| " + " | ".join(["---"] * len(cols)) + " |")
                for _, row in preview.iterrows():
                    vals = [str(row.get(c, ""))[:20] for c in cols]
                    lines.append("| " + " | ".join(vals) + " |")
                lines.append("")

            if goodcase_df is not None and not goodcase_df.empty:
                lines.append(f"### 优质案例（Good Case）\n")
                lines.append(f"> 共 {len(goodcase_df)} 条\n")
                preview = goodcase_df.head(10)
                cols = list(preview.columns)[:6]
                lines.append("| " + " | ".join(cols) + " |")
                lines.append("| " + " | ".join(["---"] * len(cols)) + " |")
                for _, row in preview.iterrows():
                    vals = [str(row.get(c, ""))[:20] for c in cols]
                    lines.append("| " + " | ".join(vals) + " |")
                lines.append("")

        return "\n".join(lines)
